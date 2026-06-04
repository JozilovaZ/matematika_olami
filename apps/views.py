from functools import wraps
from datetime import timedelta

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.models import User
from django.contrib import messages
from django.db.models import Avg, Count, Sum
from django.db.models.functions import TruncDate
from django.utils import timezone
from .models import (
    Kategoriya, Mavzu, Dars, Topshiriq, Savol, Musobaqa, MusobaqaSavol,
    MusobaqaNatija, ReytingOyinchi,
    KuchsizSoha, HaftalikFaoliyat, Modul, Yutuq, KunlikSovrin,
    OtaOnaMaslahat, UmumiyStatistika, Profil, OquvSozlama,
    BildirishnomaSozlamasi, Tema, XavfsizlikSozlamasi, DarajaSozlamasi,
    SongiFaoliyat, TestNatija,
)


def oqituvchi_talab(view_func):
    """Faqat o'qituvchi roli bo'lgan foydalanuvchilar uchun dekorator."""
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('kirish')
        if not hasattr(request.user, 'profil') or request.user.profil.rol != 'oqituvchi':
            messages.error(request, "Bu sahifaga faqat o'qituvchilar kira oladi.")
            return redirect('bosh_sahifa')
        return view_func(request, *args, **kwargs)
    return wrapper


def royxatdan_otish(request):
    if request.user.is_authenticated:
        return redirect('bosh_sahifa')

    if request.method == 'POST':
        ism = request.POST.get('ism', '').strip()
        familiya = request.POST.get('familiya', '').strip()
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        parol = request.POST.get('parol', '')
        parol2 = request.POST.get('parol2', '')
        yosh = request.POST.get('yosh', '12')
        sinf = request.POST.get('sinf', '6-sinf')
        maktab = request.POST.get('maktab', '').strip()

        xatolar = []
        if not ism:
            xatolar.append('Ismingizni kiriting')
        if not familiya:
            xatolar.append('Familiyangizni kiriting')
        if not username:
            xatolar.append('Foydalanuvchi nomini kiriting')
        if not email:
            xatolar.append('Email manzilingizni kiriting')
        if len(parol) < 6:
            xatolar.append('Parol kamida 6 ta belgidan iborat bo\'lishi kerak')
        if parol != parol2:
            xatolar.append('Parollar mos kelmadi')
        if User.objects.filter(username=username).exists():
            xatolar.append('Bu foydalanuvchi nomi allaqachon band')
        if User.objects.filter(email=email).exists():
            xatolar.append('Bu email allaqachon ro\'yxatdan o\'tgan')

        if xatolar:
            return render(request, 'royxatdan_otish.html', {
                'xatolar': xatolar,
                'forma': request.POST,
            })

        user = User.objects.create_user(
            username=username,
            email=email,
            password=parol,
            first_name=ism,
            last_name=familiya,
        )
        Profil.objects.create(
            user=user,
            ism=ism,
            familiya=familiya,
            email=email,
            rol='oquvchi',
            yosh=int(yosh) if yosh.isdigit() else 12,
            sinf=sinf,
            maktab=maktab,
        )
        login(request, user)
        messages.success(request, f'Xush kelibsiz, {ism}! Muvaffaqiyatli ro\'yxatdan o\'tdingiz.')
        return redirect('bosh_sahifa')

    return render(request, 'royxatdan_otish.html')


def kirish(request):
    if request.user.is_authenticated:
        return redirect('bosh_sahifa')

    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        parol = request.POST.get('parol', '')

        user = authenticate(request, username=username, password=parol)
        if user is not None:
            login(request, user)
            keyingi = request.GET.get('next', 'bosh_sahifa')
            return redirect(keyingi)
        else:
            return render(request, 'kirish.html', {
                'xato': 'Foydalanuvchi nomi yoki parol noto\'g\'ri',
                'username': username,
            })

    return render(request, 'kirish.html')


def chiqish(request):
    logout(request)
    return redirect('kirish')


def oqituvchi_oquvchilar(request):
    """O'qituvchi uchun: ro'yxatdan o'tgan o'quvchilar ro'yxati va statistikasi."""
    if not request.user.is_authenticated:
        return redirect('kirish')
    if not hasattr(request.user, 'profil') or request.user.profil.rol != 'oqituvchi':
        messages.error(request, "Bu sahifaga faqat o'qituvchilar kira oladi.")
        return redirect('bosh_sahifa')

    profillar = (
        Profil.objects.filter(rol='oquvchi', user__isnull=False)
        .select_related('user')
        .order_by('-user__date_joined')
    )

    oquvchilar = []
    jami_testlar = 0
    foizlar = []
    for p in profillar:
        natijalar = TestNatija.objects.filter(user=p.user)
        testlar_soni = natijalar.count()
        ortacha = int(natijalar.aggregate(a=Avg('foiz'))['a'] or 0)
        oxirgi_natija = natijalar.order_by('-sana').first()
        jami_testlar += testlar_soni
        if testlar_soni:
            foizlar.append(ortacha)
        oquvchilar.append({
            'profil': p,
            'user': p.user,
            'testlar_soni': testlar_soni,
            'ortacha': ortacha,
            'oxirgi_faollik': oxirgi_natija.sana if oxirgi_natija else p.user.last_login,
            'royxatdan': p.user.date_joined,
        })

    context = {
        'oquvchilar': oquvchilar,
        'oquvchilar_soni': len(oquvchilar),
        'jami_testlar': jami_testlar,
        'umumiy_ortacha': int(sum(foizlar) / len(foizlar)) if foizlar else 0,
        'faol_soni': sum(1 for o in oquvchilar if o['testlar_soni'] > 0),
    }
    return render(request, 'oqituvchi/oquvchilar.html', context)


def bosh_sahifa(request):
    if not request.user.is_authenticated:
        return redirect('royxatdan_otish')
    stat = UmumiyStatistika.get()
    context = {
        'darslar_soni': Dars.objects.count(),
        'test_soni': Topshiriq.objects.count(),
        'savol_soni': Savol.objects.count(),
        'stat': stat,
        'songgi_faoliyat': SongiFaoliyat.objects.filter(sahifa='bosh_sahifa'),
    }
    return render(request, 'bosh_sahifa.html', context)


def darslar(request):
    stat = UmumiyStatistika.get()
    darslar_list = Dars.objects.all()

    tugallangan = darslar_list.filter(holat='tugallangan').count()
    jarayonda = darslar_list.filter(holat='jarayonda').count()
    qolgan = darslar_list.filter(holat='qulflangan').count()

    context = {
        'darslar': darslar_list,
        'sahifalar': [1, 2, 3],
        'joriy_sahifa': 1,
        'stat': stat,
        'tugallangan': tugallangan,
        'jarayonda_soni': jarayonda,
        'qolgan': qolgan,
        'songgi_faoliyat': SongiFaoliyat.objects.filter(sahifa='darslar'),
    }
    return render(request, 'darslar.html', context)


def dars_detail(request, pk):
    import re as _re
    from django.utils.html import escape as _escape
    from django.utils.safestring import mark_safe as _mark_safe
    dars = get_object_or_404(Dars, pk=pk)
    stat = UmumiyStatistika.get()
    oldingi = Dars.objects.filter(tartib__lt=dars.tartib).order_by('-tartib').first()
    keyingi = Dars.objects.filter(tartib__gt=dars.tartib).order_by('tartib').first()
    topshiriqlar = dars.topshiriqlar.all()
    mazmun = dars.mazmun or ''
    if mazmun and not _re.search(r'<[a-zA-Z]', mazmun):
        mazmun_html = _mark_safe(
            '<p>' + _escape(mazmun).replace('\r\n', '\n').replace('\n\n', '</p><p>').replace('\n', '<br>') + '</p>'
        )
    else:
        mazmun_html = _mark_safe(mazmun)
    context = {
        'dars': dars,
        'mazmun_html': mazmun_html,
        'nazariy_html': _mark_safe(dars.nazariy) if dars.nazariy else '',
        'taqdimot_html': _mark_safe(dars.taqdimot) if dars.taqdimot else '',
        'amaliy_html': _mark_safe(dars.amaliy_uslubiy) if dars.amaliy_uslubiy else '',
        'topshiriq_matni_html': _mark_safe(dars.topshiriq_matni) if dars.topshiriq_matni else '',
        'stat': stat,
        'oldingi': oldingi,
        'keyingi': keyingi,
        'topshiriqlar': topshiriqlar,
    }
    return render(request, 'dars_detail.html', context)


def taqdimotlar(request):
    stat = UmumiyStatistika.get()
    darslar_list = Dars.objects.exclude(taqdimot='', taqdimot_fayl='')
    context = {
        'darslar': darslar_list,
        'stat': stat,
    }
    return render(request, 'taqdimotlar.html', context)


def dars_taqdimot(request, pk):
    from django.utils.safestring import mark_safe as _mark_safe
    dars = get_object_or_404(Dars, pk=pk)
    stat = UmumiyStatistika.get()
    context = {
        'dars': dars,
        'taqdimot_html': _mark_safe(dars.taqdimot) if dars.taqdimot else '',
        'stat': stat,
    }
    return render(request, 'taqdimot_detail.html', context)


def topshiriqlar(request):
    stat = UmumiyStatistika.get()
    mavzular = Mavzu.objects.prefetch_related('topshiriqlar').all()
    mavzu_topshiriqlar = [
        {'mavzu': m, 'topshiriqlar': m.topshiriqlar.all()}
        for m in mavzular
        if m.topshiriqlar.exists()
    ]
    context = {
        'mavzu_topshiriqlar': mavzu_topshiriqlar,
        'kichik_topshiriqlar': Topshiriq.objects.filter(turi='kichik', mavzu__isnull=True),
        'pastki_topshiriqlar': Topshiriq.objects.filter(turi='pastki', mavzu__isnull=True),
        'stat': stat,
        'songgi_faoliyat': SongiFaoliyat.objects.filter(sahifa='topshiriqlar'),
    }
    return render(request, 'topshiriqlar.html', context)


def topshiriq_detail(request, pk):
    import random as _random
    topshiriq = get_object_or_404(Topshiriq, pk=pk)
    stat = UmumiyStatistika.get()
    barcha = list(topshiriq.savollar_list.all())
    savollar = _random.sample(barcha, min(10, len(barcha)))
    savol_idlar = ','.join(str(s.pk) for s in savollar)
    context = {
        'topshiriq': topshiriq,
        'savollar': savollar,
        'savol_idlar': savol_idlar,
        'stat': stat,
    }
    return render(request, 'topshiriq_detail.html', context)


def topshiriq_tekshirish(request, pk):
    topshiriq = get_object_or_404(Topshiriq, pk=pk)

    if request.method != 'POST':
        return redirect('topshiriq_detail', pk=pk)

    stat = UmumiyStatistika.get()
    savol_idlar = request.POST.get('savol_idlar', '')
    if savol_idlar:
        id_list = [int(i) for i in savol_idlar.split(',') if i.strip().isdigit()]
        savollar = topshiriq.savollar_list.filter(pk__in=id_list)
    else:
        savollar = topshiriq.savollar_list.all()

    natijalar = []
    batafsil = []
    togri_soni = 0
    for savol in savollar:
        javob = request.POST.get(f'savol_{savol.pk}', '')
        togri = javob == savol.togri_javob
        if togri:
            togri_soni += 1
        natijalar.append({
            'savol': savol,
            'javob': javob,
            'togri': togri,
        })
        batafsil.append({
            'savol_id': savol.pk,
            'savol_matni': savol.savol_matni,
            'javob': javob,
            'togri_javob': savol.togri_javob,
            'togri': togri,
        })

    umumiy = savollar.count()
    foiz = int((togri_soni / umumiy) * 100) if umumiy > 0 else 0

    # Natijani bazaga saqlash
    TestNatija.objects.create(
        user=request.user if request.user.is_authenticated else None,
        topshiriq=topshiriq,
        togri_soni=togri_soni,
        umumiy_soni=umumiy,
        foiz=foiz,
        batafsil=batafsil,
    )

    # Topshiriq foizini yangilash (oxirgi natija asosida)
    topshiriq.foiz = foiz
    topshiriq.save(update_fields=['foiz'])

    context = {
        'topshiriq': topshiriq,
        'natijalar': natijalar,
        'togri_soni': togri_soni,
        'umumiy': umumiy,
        'foiz': foiz,
        'stat': stat,
    }
    return render(request, 'topshiriq_natija.html', context)


def _haftalik_faoliyat_hisoblash(natijalar_qs):
    """Oxirgi 7 kun ichida har bir kunda yechilgan testlar sonini hisoblaydi."""
    bugun = timezone.localdate()
    boshlanish = bugun - timedelta(days=6)
    kun_nomlari = ['Du', 'Se', 'Ch', 'Pa', 'Ju', 'Sh', 'Ya']

    kunlik = (
        natijalar_qs
        .filter(sana__date__gte=boshlanish)
        .annotate(kun=TruncDate('sana'))
        .values('kun')
        .annotate(soni=Count('id'), ort=Avg('foiz'))
    )
    xarita = {k['kun']: k for k in kunlik}

    natija = []
    maks = 1
    for i in range(7):
        sana = boshlanish + timedelta(days=i)
        ma = xarita.get(sana)
        soni = ma['soni'] if ma else 0
        ort = int(ma['ort']) if ma and ma['ort'] is not None else 0
        if soni > maks:
            maks = soni
        natija.append({
            'kun': kun_nomlari[sana.weekday()],
            'sana': sana,
            'testlar': soni,
            'ortacha': ort,
        })

    for d in natija:
        d['foiz'] = int((d['testlar'] / maks) * 100) if maks else 0
    return natija


def _ketma_ket_kun_hisoblash(natijalar_qs):
    """TestNatija asosida ketma-ket faol kunlar sonini topadi (bugungacha)."""
    kunlar = set(
        natijalar_qs.annotate(kun=TruncDate('sana'))
        .values_list('kun', flat=True)
        .distinct()
    )
    if not kunlar:
        return 0
    bugun = timezone.localdate()
    hisob = 0
    joriy = bugun
    if joriy not in kunlar:
        joriy = bugun - timedelta(days=1)
        if joriy not in kunlar:
            return 0
    while joriy in kunlar:
        hisob += 1
        joriy -= timedelta(days=1)
    return hisob


def _modullar_hisoblash():
    """Kategoriyalar bo'yicha o'rtacha test foizini hisoblaydi."""
    ranglar = [
        ('#e8f5e9', '#4caf50'),
        ('#fff3e0', '#ff9800'),
        ('#e3f2fd', '#2196f3'),
        ('#fce4ec', '#e91e63'),
        ('#f3e5f5', '#9c27b0'),
        ('#e0f7fa', '#00bcd4'),
    ]
    ikonlar = ['🌍', '🗺️', '📍', '🛰️', '🧭', '📊']

    natija = []
    # Faqat darsi mavjud kategoriyalar ko'rsatiladi
    kategoriyalar = [k for k in Kategoriya.objects.all() if k.darslar.exists()]
    for i, kat in enumerate(kategoriyalar):
        agr = TestNatija.objects.filter(
            topshiriq__dars__kategoriya=kat
        ).aggregate(ort=Avg('foiz'), soni=Count('id'))
        ball = int(agr['ort'] or 0)
        soni = agr['soni'] or 0

        if ball >= 80:
            holat = 'Ajoyib'
        elif ball >= 50:
            holat = "O'rtacha"
        elif soni > 0:
            holat = 'Yaxshilash kerak'
        else:
            holat = 'Boshlanmagan'

        rang, progress_rang = ranglar[i % len(ranglar)]
        natija.append({
            'nomi': kat.nomi,
            'icon': ikonlar[i % len(ikonlar)],
            'rang': rang,
            'progress_rang': progress_rang,
            'ball': ball,
            'urinishlar': soni,
            'holat': holat,
        })
    return natija


def _yutuqlar_hisoblash(natijalar_qs, ketma_ket, ortacha):
    """Foydalanuvchi yutuqlarini dinamik ravishda aniqlaydi."""
    jami = natijalar_qs.count()
    ajoyib_soni = natijalar_qs.filter(foiz__gte=80).count()
    mukammal_soni = natijalar_qs.filter(foiz=100).count()

    shartlar = [
        {
            'nomi': 'Birinchi qadam',
            'tavsif': 'Birinchi testingizni yechdingiz',
            'icon': '🎯',
            'rang': '#e8f5e9',
            'bajarildi': jami >= 1,
        },
        {
            'nomi': "10 ta test yechildi",
            'tavsif': "10 ta testni muvaffaqiyatli bajardingiz",
            'icon': '🏅',
            'rang': '#fff3e0',
            'bajarildi': jami >= 10,
        },
        {
            'nomi': 'Mukammal natija',
            'tavsif': '100% natija bilan test yechdingiz',
            'icon': '⭐',
            'rang': '#fce4ec',
            'bajarildi': mukammal_soni >= 1,
        },
        {
            'nomi': "3 kun ketma-ket",
            'tavsif': '3 kun uzluksiz mashq qildingiz',
            'icon': '🔥',
            'rang': '#e3f2fd',
            'bajarildi': ketma_ket >= 3,
        },
        {
            'nomi': "7 kun ketma-ket",
            'tavsif': 'Bir hafta uzluksiz mashq qildingiz',
            'icon': '🚀',
            'rang': '#f3e5f5',
            'bajarildi': ketma_ket >= 7,
        },
        {
            'nomi': 'A\'lochi',
            'tavsif': "5 tadan ortiq testda 80%+ oldingiz",
            'icon': '🎓',
            'rang': '#e0f7fa',
            'bajarildi': ajoyib_soni >= 5,
        },
    ]
    return [s for s in shartlar if s['bajarildi']]


def statistika(request):
    stat = UmumiyStatistika.get()

    # Foydalanuvchi asosida filter
    if request.user.is_authenticated:
        barcha_natijalar = TestNatija.objects.filter(user=request.user)
    else:
        barcha_natijalar = TestNatija.objects.all()

    oxirgi_natijalar = list(barcha_natijalar[:10])
    jami_testlar = barcha_natijalar.count()
    ortacha_foiz = int(barcha_natijalar.aggregate(avg=Avg('foiz'))['avg'] or 0)
    jami_togri = barcha_natijalar.aggregate(s=Sum('togri_soni'))['s'] or 0
    jami_savollar = barcha_natijalar.aggregate(s=Sum('umumiy_soni'))['s'] or 0
    jami_notogri = max(0, jami_savollar - jami_togri)

    # Natijalar taqsimoti (ajoyib / yaxshi / yomon)
    ajoyib = barcha_natijalar.filter(foiz__gte=80).count()
    yaxshi = barcha_natijalar.filter(foiz__gte=50, foiz__lt=80).count()
    yomon = barcha_natijalar.filter(foiz__lt=50).count()

    # Topshiriq bo'yicha statistika
    topshiriq_stat = []
    for top in Topshiriq.objects.all():
        top_natijalar = barcha_natijalar.filter(topshiriq=top)
        if top_natijalar.exists():
            top_avg = top_natijalar.aggregate(avg=Avg('foiz'))['avg'] or 0
            top_count = top_natijalar.count()
            eng_yaxshi = top_natijalar.order_by('-foiz').first()
            topshiriq_stat.append({
                'topshiriq': top,
                'urinishlar': top_count,
                'ortacha': int(top_avg),
                'eng_yaxshi': eng_yaxshi.foiz if eng_yaxshi else 0,
            })

    # Alohida bloklar
    haftalik = _haftalik_faoliyat_hisoblash(barcha_natijalar)
    ketma_ket = _ketma_ket_kun_hisoblash(barcha_natijalar)
    modullar = _modullar_hisoblash()
    yutuqlar = _yutuqlar_hisoblash(barcha_natijalar, ketma_ket, ortacha_foiz)

    # Chart.js uchun ma'lumotlar
    trend_natijalar = list(reversed(oxirgi_natijalar))
    chart = {
        'haftalik_labels': [d['kun'] for d in haftalik],
        'haftalik_testlar': [d['testlar'] for d in haftalik],
        'haftalik_ortacha': [d['ortacha'] for d in haftalik],
        'modul_labels': [m['nomi'] for m in modullar],
        'modul_ballar': [m['ball'] for m in modullar],
        'modul_ranglar': [m['progress_rang'] for m in modullar],
        'trend_labels': [n.topshiriq.nomi[:14] for n in trend_natijalar],
        'trend_foizlar': [n.foiz for n in trend_natijalar],
        'taqsimot': [ajoyib, yaxshi, yomon],
        'javoblar': [int(jami_togri), int(jami_notogri)],
    }

    context = {
        'haftalik_faoliyat': haftalik,
        'modullar': modullar,
        'yutuqlar': yutuqlar,
        'ota_ona_maslahatlar': OtaOnaMaslahat.objects.all(),
        'stat': stat,
        'ketma_ket_kun': ketma_ket,
        'oxirgi_natijalar': oxirgi_natijalar,
        'jami_testlar': jami_testlar,
        'ortacha_foiz': ortacha_foiz,
        'jami_togri': jami_togri,
        'jami_savollar': jami_savollar,
        'topshiriq_stat': topshiriq_stat,
        'ajoyib_soni': ajoyib,
        'yaxshi_soni': yaxshi,
        'yomon_soni': yomon,
        'chart': chart,
    }
    return render(request, 'statistika.html', context)


def yutuqlar(request):
    stat = UmumiyStatistika.get()
    olinganlar = Yutuq.objects.filter(olingan=True)
    olinmaganlar = Yutuq.objects.filter(olingan=False)

    olingan_soni = olinganlar.count()
    umumiy_soni = Yutuq.objects.count()
    umumiy_ball = sum(y.ball for y in olinganlar)

    darajalar = DarajaSozlamasi.objects.all()
    daraja_nomi = 'Boshlang\'ich'
    daraja_icon = '🌟'
    joriy_ball = umumiy_ball
    kerakli_ball = 500
    daraja_foiz = 0

    for i, d in enumerate(darajalar):
        if umumiy_ball >= d.kerakli_ball:
            daraja_nomi = d.nomi
            daraja_icon = d.icon
            if i + 1 < len(darajalar):
                kerakli_ball = darajalar[i + 1].kerakli_ball
            else:
                kerakli_ball = d.kerakli_ball
        else:
            kerakli_ball = d.kerakli_ball
            break

    if kerakli_ball > 0:
        oldingi_ball = 0
        for d in darajalar:
            if d.kerakli_ball >= kerakli_ball:
                break
            oldingi_ball = d.kerakli_ball
        oraliq = kerakli_ball - oldingi_ball
        if oraliq > 0:
            daraja_foiz = min(100, int(((umumiy_ball - oldingi_ball) / oraliq) * 100))
        else:
            daraja_foiz = 100

    keyingi_daraja_ball = max(0, kerakli_ball - umumiy_ball)
    oxirgi_yutuqlar = olinganlar.order_by('-sana')[:3]

    context = {
        'olinganlar': olinganlar,
        'olinmaganlar': olinmaganlar,
        'kunlik_sovrinlar': KunlikSovrin.objects.all(),
        'oxirgi_yutuqlar': oxirgi_yutuqlar,
        'olingan_yutuqlar': olingan_soni,
        'umumiy_yutuqlar': umumiy_soni,
        'umumiy_ball': umumiy_ball,
        'ketma_ket_kun': stat.ketma_ket_kun,
        'daraja': daraja_nomi,
        'daraja_icon': daraja_icon,
        'daraja_foiz': daraja_foiz,
        'keyingi_daraja_ball': keyingi_daraja_ball,
        'joriy_ball': joriy_ball,
        'kerakli_ball': kerakli_ball,
        'stat': stat,
    }
    return render(request, 'yutuqlar.html', context)


def sozlamalar(request):
    stat = UmumiyStatistika.get()
    profil = None
    if request.user.is_authenticated and hasattr(request.user, 'profil'):
        profil = request.user.profil

    xabar = None
    xato = None

    if request.method == 'POST' and profil:
        action = request.POST.get('action', 'profil')
        if action == 'profil':
            profil.ism = request.POST.get('ism', profil.ism).strip()
            profil.familiya = request.POST.get('familiya', profil.familiya).strip()
            profil.email = request.POST.get('email', profil.email).strip()
            try:
                profil.yosh = int(request.POST.get('yosh', profil.yosh))
            except ValueError:
                pass
            profil.sinf = request.POST.get('sinf', profil.sinf).strip()
            profil.maktab = request.POST.get('maktab', profil.maktab).strip()
            profil.save()
            xabar = "Profil muvaffaqiyatli saqlandi."
        elif action == 'parol':
            old_pass = request.POST.get('eski_parol', '')
            new_pass = request.POST.get('yangi_parol', '')
            confirm_pass = request.POST.get('parol_tasdiqlash', '')
            if not request.user.check_password(old_pass):
                xato = "Eski parol noto'g'ri."
            elif new_pass != confirm_pass:
                xato = "Yangi parollar mos emas."
            elif len(new_pass) < 6:
                xato = "Parol kamida 6 ta belgidan iborat bo'lishi kerak."
            else:
                request.user.set_password(new_pass)
                request.user.save()
                from django.contrib.auth import update_session_auth_hash
                update_session_auth_hash(request, request.user)
                xabar = "Parol muvaffaqiyatli o'zgartirildi."

    context = {
        'profil': profil,
        'oquv_sozlamalari': OquvSozlama.objects.all(),
        'bildirishnomalar': BildirishnomaSozlamasi.objects.all(),
        'temalar': Tema.objects.all(),
        'xavfsizlik_sozlamalari': XavfsizlikSozlamasi.objects.all(),
        'stat': stat,
        'songgi_faoliyat': SongiFaoliyat.objects.filter(sahifa='sozlamalar'),
        'xabar': xabar,
        'xato': xato,
    }
    return render(request, 'sozlamalar.html', context)


def _musobaqa_top10(musobaqa):
    """Har bir foydalanuvchining eng yaxshi natijasini olish va top 10 ni qaytarish."""
    from django.db.models import Max
    natijalar = (
        MusobaqaNatija.objects
        .filter(musobaqa=musobaqa, user__isnull=False)
        .values('user__id', 'user__first_name', 'user__last_name', 'user__username')
        .annotate(eng_yaxshi=Max('foiz'))
        .order_by('-eng_yaxshi')[:10]
    )
    top10 = []
    for i, n in enumerate(natijalar):
        ism = n['user__first_name'] or n['user__username']
        familiya = n['user__last_name'] or ''
        top10.append({
            'ism': f"{ism} {familiya}".strip(),
            'foiz': n['eng_yaxshi'],
            'orni': i + 1,
        })
    return top10


def musobaqa_detail(request, pk):
    musobaqa = get_object_or_404(Musobaqa, pk=pk)
    stat = UmumiyStatistika.get()
    savollar = musobaqa.savollar.all()
    top10 = _musobaqa_top10(musobaqa)
    context = {
        'musobaqa': musobaqa,
        'savollar': savollar,
        'stat': stat,
        'top10': top10,
    }
    return render(request, 'musobaqa_detail.html', context)


def musobaqa_tekshirish(request, pk):
    musobaqa = get_object_or_404(Musobaqa, pk=pk)

    if request.method != 'POST':
        return redirect('musobaqa_detail', pk=pk)

    stat = UmumiyStatistika.get()
    savollar = musobaqa.savollar.all()

    natijalar = []
    batafsil = []
    togri_soni = 0
    for savol in savollar:
        javob = request.POST.get(f'savol_{savol.pk}', '')
        togri = javob == savol.togri_javob
        if togri:
            togri_soni += 1
        natijalar.append({
            'savol': savol,
            'javob': javob,
            'togri': togri,
        })
        batafsil.append({
            'savol_id': savol.pk,
            'savol_matni': savol.savol_matni,
            'javob': javob,
            'togri_javob': savol.togri_javob,
            'togri': togri,
        })

    umumiy = savollar.count()
    foiz = int((togri_soni / umumiy) * 100) if umumiy > 0 else 0

    MusobaqaNatija.objects.create(
        user=request.user if request.user.is_authenticated else None,
        musobaqa=musobaqa,
        togri_soni=togri_soni,
        umumiy_soni=umumiy,
        foiz=foiz,
        batafsil=batafsil,
    )

    # Qatnashchilar sonini yangilash
    musobaqa.qatnashchilar = musobaqa.natijalar.count()
    musobaqa.save(update_fields=['qatnashchilar'])

    top10 = _musobaqa_top10(musobaqa)

    context = {
        'musobaqa': musobaqa,
        'natijalar': natijalar,
        'togri_soni': togri_soni,
        'umumiy': umumiy,
        'foiz': foiz,
        'stat': stat,
        'top10': top10,
    }
    return render(request, 'musobaqa_natija.html', context)


# ===================== O'QITUVCHI AMALIYOTLARI =====================

# --- Kategoriya ---


@oqituvchi_talab
def oqituvchi_kategoriya_qoshish(request):
    if request.method == 'POST':
        from django.db.models import Max as _Max
        oxirgi = (Kategoriya.objects.aggregate(m=_Max('tartib'))['m'] or 0) + 1
        Kategoriya.objects.create(
            nomi=request.POST.get('nomi', '').strip(),
            tartib=oxirgi,
        )
        messages.success(request, 'Kategoriya qo\'shildi.')
        return redirect('bosh_sahifa')
    return render(request, 'oqituvchi/kategoriya_forma.html', {
        'sarlavha': 'Yangi kategoriya',
    })


@oqituvchi_talab
def oqituvchi_kategoriya_tahrirlash(request, pk):
    obj = get_object_or_404(Kategoriya, pk=pk)
    if request.method == 'POST':
        obj.nomi = request.POST.get('nomi', '').strip()
        obj.save()
        messages.success(request, 'Kategoriya yangilandi.')
        return redirect('bosh_sahifa')
    return render(request, 'oqituvchi/kategoriya_forma.html', {
        'sarlavha': f'"{obj.nomi}" ni tahrirlash',
        'obj': obj,
    })


@oqituvchi_talab
def oqituvchi_kategoriya_ochirish(request, pk):
    obj = get_object_or_404(Kategoriya, pk=pk)
    if request.method == 'POST':
        obj.delete()
        messages.success(request, 'Kategoriya o\'chirildi.')
    return redirect('bosh_sahifa')


# --- Mavzu ---
@oqituvchi_talab
def oqituvchi_mavzu_qoshish(request):
    if request.method == 'POST':
        from django.db.models import Max as _Max
        oxirgi = (Mavzu.objects.aggregate(m=_Max('tartib'))['m'] or 0) + 1
        Mavzu.objects.create(
            raqam=request.POST.get('raqam', '').strip(),
            nomi=request.POST.get('nomi', '').strip(),
            tavsif=request.POST.get('tavsif', '').strip(),
            rang=request.POST.get('rang', '#e8f5e9'),
            icon=request.POST.get('icon', '').strip(),
            tartib=oxirgi,
        )
        messages.success(request, 'Mavzu qo\'shildi.')
        return redirect('bosh_sahifa')
    return render(request, 'oqituvchi/mavzu_forma.html', {
        'sarlavha': 'Yangi mavzu',
    })


@oqituvchi_talab
def oqituvchi_mavzu_tahrirlash(request, pk):
    obj = get_object_or_404(Mavzu, pk=pk)
    if request.method == 'POST':
        obj.raqam = request.POST.get('raqam', '').strip()
        obj.nomi = request.POST.get('nomi', '').strip()
        obj.tavsif = request.POST.get('tavsif', '').strip()
        obj.rang = request.POST.get('rang', '#e8f5e9')
        obj.icon = request.POST.get('icon', '').strip()
        obj.save()
        messages.success(request, 'Mavzu yangilandi.')
        return redirect('bosh_sahifa')
    return render(request, 'oqituvchi/mavzu_forma.html', {
        'sarlavha': f'"{obj.nomi}" ni tahrirlash',
        'obj': obj,
    })


@oqituvchi_talab
def oqituvchi_mavzu_ochirish(request, pk):
    obj = get_object_or_404(Mavzu, pk=pk)
    if request.method == 'POST':
        obj.delete()
        messages.success(request, 'Mavzu o\'chirildi.')
    return redirect('bosh_sahifa')


# --- Dars ---
@oqituvchi_talab
def oqituvchi_dars_qoshish(request):
    if request.method == 'POST':
        kat_pk = request.POST.get('kategoriya')
        kat = Kategoriya.objects.filter(pk=kat_pk).first() if kat_pk else None
        from django.db.models import Max as _Max
        oxirgi_tartib = (Dars.objects.aggregate(m=_Max('tartib'))['m'] or 0) + 1
        Dars.objects.create(
            raqam=request.POST.get('raqam', '').strip(),
            nomi=request.POST.get('nomi', '').strip(),
            tavsif=request.POST.get('tavsif', '').strip(),
            icon=request.POST.get('icon', '').strip(),
            rang=request.POST.get('rang', '#e8f5e9'),
            davomiylik=request.POST.get('davomiylik', '').strip(),
            holat=request.POST.get('holat', 'qulflangan'),
            kategoriya=kat,
            mazmun=request.POST.get('mazmun', '').strip(),
            tartib=oxirgi_tartib,
        )
        messages.success(request, 'Dars qo\'shildi.')
        return redirect('darslar')
    return render(request, 'oqituvchi/dars_forma.html', {
        'sarlavha': 'Yangi dars',
        'kategoriyalar': Kategoriya.objects.all(),
    })


@oqituvchi_talab
def oqituvchi_dars_tahrirlash(request, pk):
    obj = get_object_or_404(Dars, pk=pk)
    if request.method == 'POST':
        kat_pk = request.POST.get('kategoriya')
        obj.raqam = request.POST.get('raqam', '').strip()
        obj.nomi = request.POST.get('nomi', '').strip()
        obj.tavsif = request.POST.get('tavsif', '').strip()
        obj.icon = request.POST.get('icon', '').strip()
        obj.rang = request.POST.get('rang', '#e8f5e9')
        obj.davomiylik = request.POST.get('davomiylik', '').strip()
        obj.holat = request.POST.get('holat', 'qulflangan')
        obj.kategoriya = Kategoriya.objects.filter(pk=kat_pk).first() if kat_pk else None
        obj.mazmun = request.POST.get('mazmun', '').strip()
        obj.save()
        messages.success(request, 'Dars yangilandi.')
        return redirect('darslar')
    return render(request, 'oqituvchi/dars_forma.html', {
        'sarlavha': f'"{obj.nomi}" ni tahrirlash',
        'obj': obj,
        'kategoriyalar': Kategoriya.objects.all(),
    })


@oqituvchi_talab
def oqituvchi_dars_ochirish(request, pk):
    obj = get_object_or_404(Dars, pk=pk)
    if request.method == 'POST':
        obj.delete()
        messages.success(request, 'Dars o\'chirildi.')
    return redirect('darslar')


# --- Topshiriq ---
@oqituvchi_talab
def oqituvchi_topshiriq_qoshish(request):
    if request.method == 'POST':
        from django.db.models import Max as _Max
        oxirgi = (Topshiriq.objects.aggregate(m=_Max('tartib'))['m'] or 0) + 1
        topshiriq = Topshiriq.objects.create(
            nomi=request.POST.get('nomi', '').strip(),
            tavsif=request.POST.get('tavsif', '').strip(),
            icon=request.POST.get('icon', 'fas fa-tasks'),
            rang=request.POST.get('rang', '#e8f5e9'),
            turi=request.POST.get('turi', 'kichik'),
            rasm_url=request.POST.get('rasm_url', ''),
            tartib=oxirgi,
        )
        messages.success(request, f'"{topshiriq.nomi}" topshirig\'i qo\'shildi.')
        return redirect('oqituvchi_topshiriq_tahrirlash', pk=topshiriq.pk)
    return render(request, 'oqituvchi/topshiriq_forma.html', {
        'sarlavha': "Yangi topshiriq qo'shish",
    })


@oqituvchi_talab
def oqituvchi_topshiriq_tahrirlash(request, pk):
    topshiriq = get_object_or_404(Topshiriq, pk=pk)
    if request.method == 'POST':
        topshiriq.nomi = request.POST.get('nomi', '').strip()
        topshiriq.tavsif = request.POST.get('tavsif', '').strip()
        topshiriq.icon = request.POST.get('icon', 'fas fa-tasks')
        topshiriq.rang = request.POST.get('rang', '#e8f5e9')
        topshiriq.turi = request.POST.get('turi', 'kichik')
        topshiriq.rasm_url = request.POST.get('rasm_url', '')
        topshiriq.save()
        messages.success(request, f'"{topshiriq.nomi}" topshirig\'i yangilandi.')
        return redirect('oqituvchi_topshiriq_tahrirlash', pk=topshiriq.pk)
    savollar = topshiriq.savollar_list.all()
    return render(request, 'oqituvchi/topshiriq_forma.html', {
        'sarlavha': f'"{topshiriq.nomi}" ni tahrirlash',
        'topshiriq': topshiriq,
        'savollar': savollar,
    })


@oqituvchi_talab
def oqituvchi_topshiriq_ochirish(request, pk):
    topshiriq = get_object_or_404(Topshiriq, pk=pk)
    if request.method == 'POST':
        topshiriq.delete()
        messages.success(request, 'Topshiriq o\'chirildi.')
    return redirect('topshiriqlar')


# --- Savol (Topshiriq savollari) ---
@oqituvchi_talab
def oqituvchi_savol_qoshish(request, topshiriq_pk):
    topshiriq = get_object_or_404(Topshiriq, pk=topshiriq_pk)
    if request.method == 'POST':
        oxirgi_tartib = topshiriq.savollar_list.count()
        Savol.objects.create(
            topshiriq=topshiriq,
            savol_matni=request.POST.get('savol_matni', '').strip(),
            variant_a=request.POST.get('variant_a', '').strip(),
            variant_b=request.POST.get('variant_b', '').strip(),
            variant_c=request.POST.get('variant_c', '').strip(),
            variant_d=request.POST.get('variant_d', '').strip(),
            togri_javob=request.POST.get('togri_javob', 'a'),
            izoh=request.POST.get('izoh', '').strip(),
            tartib=oxirgi_tartib,
        )
        topshiriq.savollar = topshiriq.savollar_list.count()
        topshiriq.save(update_fields=['savollar'])
        messages.success(request, 'Savol qo\'shildi.')
        return redirect('oqituvchi_topshiriq_tahrirlash', pk=topshiriq.pk)
    return render(request, 'oqituvchi/savol_forma.html', {
        'sarlavha': f'"{topshiriq.nomi}" ga savol qo\'shish',
        'topshiriq': topshiriq,
    })


@oqituvchi_talab
def oqituvchi_savol_tahrirlash(request, pk):
    savol = get_object_or_404(Savol, pk=pk)
    if request.method == 'POST':
        savol.savol_matni = request.POST.get('savol_matni', '').strip()
        savol.variant_a = request.POST.get('variant_a', '').strip()
        savol.variant_b = request.POST.get('variant_b', '').strip()
        savol.variant_c = request.POST.get('variant_c', '').strip()
        savol.variant_d = request.POST.get('variant_d', '').strip()
        savol.togri_javob = request.POST.get('togri_javob', 'a')
        savol.izoh = request.POST.get('izoh', '').strip()
        savol.save()
        messages.success(request, 'Savol yangilandi.')
        return redirect('oqituvchi_topshiriq_tahrirlash', pk=savol.topshiriq.pk)
    return render(request, 'oqituvchi/savol_forma.html', {
        'sarlavha': 'Savolni tahrirlash',
        'topshiriq': savol.topshiriq,
        'savol': savol,
    })


@oqituvchi_talab
def oqituvchi_savol_ochirish(request, pk):
    savol = get_object_or_404(Savol, pk=pk)
    topshiriq = savol.topshiriq
    if request.method == 'POST':
        savol.delete()
        topshiriq.savollar = topshiriq.savollar_list.count()
        topshiriq.save(update_fields=['savollar'])
        messages.success(request, 'Savol o\'chirildi.')
    return redirect('oqituvchi_topshiriq_tahrirlash', pk=topshiriq.pk)


# --- Musobaqa ---
@oqituvchi_talab
def oqituvchi_musobaqa_qoshish(request):
    if request.method == 'POST':
        musobaqa = Musobaqa.objects.create(
            nomi=request.POST.get('nomi', '').strip(),
            tavsif=request.POST.get('tavsif', '').strip(),
            icon=request.POST.get('icon', '').strip(),
            rang=request.POST.get('rang', '#e8f5e9'),
            daraja=request.POST.get('daraja', "O'rta"),
            tartib=int(request.POST.get('tartib', 0)),
        )
        messages.success(request, f'"{musobaqa.nomi}" musobaqasi qo\'shildi.')
        return redirect('oqituvchi_musobaqa_tahrirlash', pk=musobaqa.pk)
    return render(request, 'oqituvchi/musobaqa_forma.html', {
        'sarlavha': 'Yangi musobaqa',
    })


@oqituvchi_talab
def oqituvchi_musobaqa_tahrirlash(request, pk):
    obj = get_object_or_404(Musobaqa, pk=pk)
    if request.method == 'POST':
        obj.nomi = request.POST.get('nomi', '').strip()
        obj.tavsif = request.POST.get('tavsif', '').strip()
        obj.icon = request.POST.get('icon', '').strip()
        obj.rang = request.POST.get('rang', '#e8f5e9')
        obj.daraja = request.POST.get('daraja', "O'rta")
        obj.tartib = int(request.POST.get('tartib', 0))
        obj.save()
        messages.success(request, f'"{obj.nomi}" musobaqasi yangilandi.')
        return redirect('oqituvchi_musobaqa_tahrirlash', pk=obj.pk)
    savollar = obj.savollar.all()
    return render(request, 'oqituvchi/musobaqa_forma.html', {
        'sarlavha': f'"{obj.nomi}" ni tahrirlash',
        'obj': obj,
        'savollar': savollar,
    })


@oqituvchi_talab
def oqituvchi_musobaqa_ochirish(request, pk):
    obj = get_object_or_404(Musobaqa, pk=pk)
    if request.method == 'POST':
        obj.delete()
        messages.success(request, 'Musobaqa o\'chirildi.')
    return redirect('bosh_sahifa')


# --- Musobaqa savollari ---
@oqituvchi_talab
def oqituvchi_musobaqa_savol_qoshish(request, musobaqa_pk):
    musobaqa = get_object_or_404(Musobaqa, pk=musobaqa_pk)
    if request.method == 'POST':
        oxirgi_tartib = musobaqa.savollar.count()
        MusobaqaSavol.objects.create(
            musobaqa=musobaqa,
            savol_matni=request.POST.get('savol_matni', '').strip(),
            variant_a=request.POST.get('variant_a', '').strip(),
            variant_b=request.POST.get('variant_b', '').strip(),
            variant_c=request.POST.get('variant_c', '').strip(),
            variant_d=request.POST.get('variant_d', '').strip(),
            togri_javob=request.POST.get('togri_javob', 'a'),
            izoh=request.POST.get('izoh', '').strip(),
            tartib=oxirgi_tartib,
        )
        messages.success(request, 'Savol qo\'shildi.')
        return redirect('oqituvchi_musobaqa_tahrirlash', pk=musobaqa.pk)
    return render(request, 'oqituvchi/savol_forma.html', {
        'sarlavha': f'"{musobaqa.nomi}" ga savol qo\'shish',
        'musobaqa': musobaqa,
    })


@oqituvchi_talab
def oqituvchi_musobaqa_savol_tahrirlash(request, pk):
    savol = get_object_or_404(MusobaqaSavol, pk=pk)
    if request.method == 'POST':
        savol.savol_matni = request.POST.get('savol_matni', '').strip()
        savol.variant_a = request.POST.get('variant_a', '').strip()
        savol.variant_b = request.POST.get('variant_b', '').strip()
        savol.variant_c = request.POST.get('variant_c', '').strip()
        savol.variant_d = request.POST.get('variant_d', '').strip()
        savol.togri_javob = request.POST.get('togri_javob', 'a')
        savol.izoh = request.POST.get('izoh', '').strip()
        savol.save()
        messages.success(request, 'Savol yangilandi.')
        return redirect('oqituvchi_musobaqa_tahrirlash', pk=savol.musobaqa.pk)
    return render(request, 'oqituvchi/savol_forma.html', {
        'sarlavha': 'Savolni tahrirlash',
        'musobaqa': savol.musobaqa,
        'savol': savol,
    })


@oqituvchi_talab
def oqituvchi_musobaqa_savol_ochirish(request, pk):
    savol = get_object_or_404(MusobaqaSavol, pk=pk)
    musobaqa = savol.musobaqa
    if request.method == 'POST':
        savol.delete()
        messages.success(request, 'Savol o\'chirildi.')
    return redirect('oqituvchi_musobaqa_tahrirlash', pk=musobaqa.pk)


# --- Word fayldan savollarni import qilish ---

def _parse_inline(line):
    """
    Bitta qatordagi savol formatini parse qiladi:
    "1. Savol matni? A) 6 B) 7 C) 8 D) 9 Javob: C"
    """
    import re
    # Savol matnini ajratib olamiz (raqam + nuqta/qavs ixtiyoriy)
    m = re.match(
        r'^(?:\d+[.)]\s*)?(.+?)\s+'
        r'[Aa][.)]\s*(.+?)\s+'
        r'[Bb][.)]\s*(.+?)\s+'
        r'[Cc][.)]\s*(.+?)\s+'
        r'[Dd][.)]\s*(.+?)\s+'
        r'[Jj]avob\s*[:=]\s*([AaBbCcDd])'
        r'(?:\s+[Ii]zoh\s*[:=]\s*(.+))?$',
        line.strip()
    )
    if m:
        return {
            'savol_matni': m.group(1).strip(),
            'variant_a':   m.group(2).strip(),
            'variant_b':   m.group(3).strip(),
            'variant_c':   m.group(4).strip(),
            'variant_d':   m.group(5).strip(),
            'togri_javob': m.group(6).lower(),
            'izoh':        (m.group(7) or '').strip(),
        }
    return None


def _word_savollarni_parse(doc):
    """
    Word fayldan savollarni o'qiydi.

    Uch xil formatni qo'llab-quvvatlaydi:

    FORMAT 1 — har bir element alohida qatorda:
      1. Savol matni?
      A) variant
      B) variant
      C) variant
      D) variant
      Javob: B

    FORMAT 2 — hammasi bitta qatorda:
      1. Savol matni? A) 6 B) 7 C) 8 D) 9 Javob: C

    FORMAT 3 — Word numbered list (raqam yo'q):
      Savol matni?
      A) variant  ...  Javob: B
    """
    import re

    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for sub in cell.text.splitlines():
                    s = sub.strip()
                    if s:
                        lines.append(s)

    savollar = []

    # --- FORMAT 2: bitta qatorda ---
    inline_found = False
    for line in lines:
        parsed = _parse_inline(line)
        if parsed:
            inline_found = True
            if not savollar or savollar[-1]['savol_matni'] != parsed['savol_matni']:
                savollar.append(parsed)

    if inline_found:
        return savollar

    # --- FORMAT 1 & 3: alohida qatorlar ---
    i = 0
    while i < len(lines):
        line = lines[i]
        a_match = re.match(r'^[Aa][.)]\s*(.+)', line)
        if a_match and i > 0:
            prev = lines[i - 1]
            savol_matni = re.sub(r'^\d+[.)]\s*', '', prev).strip()
            variant = {
                'a': a_match.group(1).strip(),
                'b': '', 'c': '', 'd': '',
                'togri': 'a', 'izoh': ''
            }
            i += 1
            while i < len(lines):
                l = lines[i].strip()
                v = re.match(r'^([BbCcDd])[.)]\s*(.+)', l)
                if v:
                    variant[v.group(1).lower()] = v.group(2).strip()
                    i += 1
                    continue
                j = re.match(r'^[Jj]avob\s*[:=]\s*([AaBbCcDd])', l)
                if j:
                    variant['togri'] = j.group(1).lower()
                    i += 1
                    continue
                iz = re.match(r'^[Ii]zoh\s*[:=]\s*(.+)', l)
                if iz:
                    variant['izoh'] = iz.group(1).strip()
                    i += 1
                    continue
                break
            if savol_matni and variant['a'] and variant['b']:
                if not savollar or savollar[-1]['savol_matni'] != savol_matni:
                    savollar.append({
                        'savol_matni': savol_matni,
                        'variant_a': variant['a'],
                        'variant_b': variant['b'],
                        'variant_c': variant['c'],
                        'variant_d': variant['d'],
                        'togri_javob': variant['togri'],
                        'izoh': variant['izoh'],
                    })
        else:
            i += 1

    return savollar


@oqituvchi_talab
def oqituvchi_savollar_import(request, topshiriq_pk):
    from docx import Document
    topshiriq = get_object_or_404(Topshiriq, pk=topshiriq_pk)

    if request.method == 'POST':
        fayl = request.FILES.get('word_fayl')
        if not fayl:
            messages.error(request, 'Fayl tanlanmadi.')
            return redirect('oqituvchi_savollar_import', topshiriq_pk=topshiriq_pk)

        if not fayl.name.endswith(('.docx', '.doc')):
            messages.error(request, 'Faqat .docx formatdagi fayl qabul qilinadi.')
            return redirect('oqituvchi_savollar_import', topshiriq_pk=topshiriq_pk)

        try:
            doc = Document(fayl)
            # Debug: fayldan o'qilgan satrlarni ko'rsatamiz
            raw_lines = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    raw_lines.append(t)
            savollar = _word_savollarni_parse(doc)
        except Exception as e:
            messages.error(request, f'Faylni o\'qishda xato: {e}')
            return redirect('oqituvchi_savollar_import', topshiriq_pk=topshiriq_pk)

        # Ko'rib chiqish yoki saqlash
        if 'preview' in request.POST or not savollar:
            return render(request, 'oqituvchi/savollar_import.html', {
                'topshiriq': topshiriq,
                'raw_lines': raw_lines,
                'savollar_preview': savollar,
                'xato': not savollar,
            })

        oxirgi_tartib = topshiriq.savollar_list.count()
        for idx, s in enumerate(savollar):
            Savol.objects.create(
                topshiriq=topshiriq,
                savol_matni=s['savol_matni'],
                variant_a=s['variant_a'],
                variant_b=s['variant_b'],
                variant_c=s['variant_c'],
                variant_d=s['variant_d'],
                togri_javob=s['togri_javob'],
                izoh=s['izoh'],
                tartib=oxirgi_tartib + idx,
            )
        topshiriq.savollar = topshiriq.savollar_list.count()
        topshiriq.save(update_fields=['savollar'])
        messages.success(request, f'{len(savollar)} ta savol muvaffaqiyatli yuklandi!')
        return redirect('oqituvchi_topshiriq_tahrirlash', pk=topshiriq_pk)

    return render(request, 'oqituvchi/savollar_import.html', {
        'topshiriq': topshiriq,
    })
